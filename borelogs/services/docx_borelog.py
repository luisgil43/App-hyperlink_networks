# borelogs/services/docx_borelog.py
from __future__ import annotations

import re
from io import BytesIO
from typing import Any, Dict, Optional, Tuple

from docx import Document


def _norm_text(s: str) -> str:
    return re.sub(r"\s+", " ", (s or "").strip())


def _cell_text(cell) -> str:
    parts = []
    for p in cell.paragraphs:
        t = (p.text or "").strip()
        if t:
            parts.append(t)
    return _norm_text(" ".join(parts))


def _set_cell_text_preserve(cell, text: str) -> None:
    """
    Escribe en celda preservando formato lo máximo posible.
    """
    text = "" if text is None else str(text)

    if not cell.paragraphs:
        cell.add_paragraph(text)
        return

    # Limpia runs sin romper estilos
    for p in cell.paragraphs:
        for r in p.runs:
            r.text = ""

    p0 = cell.paragraphs[0]
    if p0.runs:
        p0.runs[0].text = text
    else:
        p0.add_run(text)


def _set_paragraph_text_preserve(paragraph, new_text: str) -> None:
    """
    Fallback: reemplaza texto de un párrafo sin destruir estilos.
    """
    for r in paragraph.runs:
        r.text = ""
    if paragraph.runs:
        paragraph.runs[0].text = new_text
    else:
        paragraph.add_run(new_text)


def build_rod_cell_map(doc_bytes: bytes, max_rods: int = 50) -> Dict[str, Dict[str, Tuple[int, int, int]]]:
    """
    Encuentra Rod # (1..max_rods) y asume Depth/Pitch/Station a la derecha en la misma fila.
    (Esto está en el body y ya te funcionó 50/50.)
    """
    doc = Document(BytesIO(doc_bytes))
    cell_map: Dict[str, Dict[str, Tuple[int, int, int]]] = {}
    rod_re = re.compile(r"^\d+$")

    for t_idx, table in enumerate(doc.tables):
        for r_idx, row in enumerate(table.rows):
            cells = row.cells
            for c_idx, cell in enumerate(cells):
                val = _cell_text(cell)
                if not val or not rod_re.match(val):
                    continue

                rod_n = int(val)
                if rod_n < 1 or rod_n > max_rods:
                    continue

                if c_idx + 3 >= len(cells):
                    continue

                cell_map[str(rod_n)] = {
                    "depth": (t_idx, r_idx, c_idx + 1),
                    "pitch": (t_idx, r_idx, c_idx + 2),
                    "station": (t_idx, r_idx, c_idx + 3),
                }

    return cell_map


def _iter_tables_with_scope(doc: Document):
    """
    Itera todas las tablas del documento incluyendo:
    - body (doc.tables)
    - header (doc.sections[i].header.tables)

    Retorna tuplas:
      ("body", None, table_idx, table)
      ("header", section_idx, table_idx, table)
    """
    for t_idx, t in enumerate(doc.tables):
        yield ("body", None, t_idx, t)

    for s_idx, section in enumerate(doc.sections):
        hdr = section.header
        for t_idx, t in enumerate(hdr.tables):
            yield ("header", s_idx, t_idx, t)


def build_header_cell_map(doc_bytes: bytes) -> Dict[str, Dict[str, int]]:
    """
    Encuentra las celdas 'cuadro' del header buscando labels en tablas.
    Devuelve un mapa con localización completa (body/header).
    Formato:
      {
        "rod_length": {"scope":"header","section":0,"table":0,"row":X,"col":Y},
        ...
      }
    """
    doc = Document(BytesIO(doc_bytes))

    labels = {
        "rod_length": "rod length",
        "driller_name": "driller name",
        "vendor_name": "vendor name",
        "project_name": "project name",
    }

    out: Dict[str, Dict[str, int]] = {}

    for scope, s_idx, t_idx, table in _iter_tables_with_scope(doc):
        for r_idx, row in enumerate(table.rows):
            cells = row.cells
            for c_idx, cell in enumerate(cells):
                txt = _cell_text(cell).lower()

                for key, needle in labels.items():
                    if key in out:
                        continue

                    # Si la celda contiene "Rod Length:" etc, el valor está en la celda de al lado
                    if needle in txt:
                        if c_idx + 1 < len(cells):
                            out[key] = {
                                "scope": 0 if scope == "body" else 1,   # 0=body, 1=header (JSON simple)
                                "section": -1 if s_idx is None else int(s_idx),
                                "table": int(t_idx),
                                "row": int(r_idx),
                                "col": int(c_idx + 1),
                            }

    return out


def _get_cell_by_locator(doc: Document, loc: Any):
    """
    Soporta 2 formatos:
    - tuple/list (t,r,c) -> body doc.tables
    - dict con scope/section/table/row/col -> body/header
    """
    # Formato antiguo (tuple)
    if isinstance(loc, (list, tuple)) and len(loc) == 3:
        t_idx, r_idx, c_idx = loc
        return doc.tables[int(t_idx)].rows[int(r_idx)].cells[int(c_idx)]

    # Formato nuevo (dict)
    if isinstance(loc, dict):
        scope = int(loc.get("scope", 0))  # 0 body, 1 header
        section = int(loc.get("section", -1))
        t_idx = int(loc["table"])
        r_idx = int(loc["row"])
        c_idx = int(loc["col"])

        if scope == 0:
            return doc.tables[t_idx].rows[r_idx].cells[c_idx]

        # header
        if section < 0:
            section = 0
        return doc.sections[section].header.tables[t_idx].rows[r_idx].cells[c_idx]

    raise ValueError("Invalid locator format for header cell.")


def _fill_header_fallback_paragraphs(doc: Document, header_values: Dict[str, str]) -> None:
    """
    Fallback (si el template no trae tablas para el header).
    Ojo: puede quedar corrido según tabs del documento.
    """
    targets = {
        "rod length": ("rod_length", "Rod Length:"),
        "driller name": ("driller_name", "Driller Name:"),
        "vendor name": ("vendor_name", "Vendor Name:"),
        "project name": ("project_name", "Project Name:"),
    }

    for p in doc.paragraphs:
        raw = p.text or ""
        low = raw.lower().strip()

        for needle, (key, _label) in targets.items():
            if needle in low and ":" in raw:
                value = header_values.get(key, "") or ""
                before, sep, after = raw.partition(":")
                # Conserva tabs/espacios existentes
                if after.strip() == "":
                    after = "\t"
                new_text = before + ":" + after + value
                _set_paragraph_text_preserve(p, new_text)
                break


def render_borelog_docx(
    template_bytes: bytes,
    rod_cell_map: Dict[str, Dict[str, Tuple[int, int, int]]],
    header_cell_map: Optional[Dict[str, Any]],
    header_values: Dict[str, str],
    rod_values: Dict[int, Dict[str, str]],
) -> bytes:
    """
    Rellena template con:
    - header en celdas (ideal: dentro de los cuadros)
    - fallback por párrafos si no detecta header en tablas
    - rods (depth/pitch/station) en tabla
    """
    doc = Document(BytesIO(template_bytes))

    # ✅ 1) HEADER dentro de cajas (tablas body/header)
    header_filled = False
    if header_cell_map:
        for key, loc in header_cell_map.items():
            if key not in header_values:
                continue
            try:
                cell = _get_cell_by_locator(doc, loc)
                _set_cell_text_preserve(cell, header_values.get(key, ""))
                header_filled = True
            except Exception:
                # si alguna coordenada falla, seguimos con el resto
                pass

    # ✅ 2) Fallback por párrafos solo si no se pudo llenar por celdas
    if not header_filled:
        _fill_header_fallback_paragraphs(doc, header_values)

    # ✅ 3) Rods en tabla (body)
    for rod_n, vals in rod_values.items():
        key = str(rod_n)
        if key not in rod_cell_map:
            continue

        mapping = rod_cell_map[key]
        for field in ("depth", "pitch", "station"):
            if field not in mapping:
                continue
            t_idx, r_idx, c_idx = mapping[field]
            try:
                cell = doc.tables[int(t_idx)].rows[int(r_idx)].cells[int(c_idx)]
            except Exception:
                continue
            _set_cell_text_preserve(cell, vals.get(field, ""))

    out = BytesIO()
    doc.save(out)
    return out.getvalue()